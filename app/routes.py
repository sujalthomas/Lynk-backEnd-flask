import time
import token
from flask import Flask, request, jsonify, url_for, render_template
from . import UPLOAD_FOLDER, app, db, serializer, mail
from .models import Resume, Role, User
from .utils import is_api_key_valid, convert_to_txt
import openai
import logging
from docx import Document
import os
from flask_mail import Mail as FlaskMail, Message
from werkzeug.security import generate_password_hash as encrypt_password
from werkzeug.security import check_password_hash as verify_password
from flask_security import (
    roles_required,
    login_required,
)
from werkzeug.security import generate_password_hash
from flask import send_file
import io
from flask import redirect


# Routes ###############
# Define routes #######
# api key verification
@app.route("/apiverify", methods=["POST"])
def apiverify():
    data = request.get_json()
    api_key = data.get("apiKey")

    try:
        if is_api_key_valid(api_key):
            token = serializer.dumps({"user": "YOUR_USER_IDENTIFIER"})
            return jsonify(success=True, token=token)
        else:
            raise Exception("Invalid API key")
    except Exception as e:
        logging.error(str(e))
        return jsonify(success=False, message="Invalid API key"), 401

user_counters = {}

# cover letter generation
@app.route("/cover-letter", methods=["POST"])
def listen():
    token = request.headers.get("Authorization")
    try:
        data = serializer.loads(token, salt="password-reset", max_age=9600)
    except Exception as e:
        print(f"Token deserialization error: {e}")
        return jsonify(success=False, message="Invalid or expired token"), 401


    data = request.get_json()
    api_key = data.get("apiKey")
    openai.api_key = api_key

    # request body data from the frontend and assign it as user_id


    company_name = data.get("Company-name", "")
    job_listing = data.get("Job-Listing", "")
    recruiter = data.get("Recruiter", "")
    date = data.get("Date", "")
    user_id = data.get("user_id")

    decoded_data = serializer.loads(user_id, salt="password-reset", max_age=9600)
    
    user_id_decode = decoded_data["user"]
    print("User ID:", user_id_decode)

        # Construct the path to the user's folder and the resume inside
    user_folder_path = os.path.join(UPLOAD_FOLDER, str(user_id_decode))
    resume_path = os.path.join(user_folder_path, "current_resume.txt")

    try:
        with open(resume_path, "r") as file:
            resume = file.read()
    except FileNotFoundError:
        return jsonify(success=False, message="Resume file not found."), 404

    try:
        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system",
                    "content": "You are a helpful assistant that will craft a tailored cover letter using a resume and a job listing. Your goal is to create a compelling cover letter that showcases the candidate's skills and experiences, aligning them with the job's requirements.",
                },
                {
                    "role": "user",
                    "content": f"Using this resume, {resume}, and this job listing, {job_listing}, craft a cover letter that doesn't include addresses but highlights the candidate's fit for the role. Ensure it includes the candidate's name, email, phone number, and LinkedIn profile. Also, only include the company name {company_name}, followed by recruiter's name {recruiter} and today's date {date}. No place holder text is allowed, if recruiters name is not found use 'Dear Hiring Manager.' ",
                },
            ],
            temperature=1.3,
            top_p=0.9,
            max_tokens=700,
            frequency_penalty=0.5,
            presence_penalty=0.5,
        )
    except openai.error.OpenAIError as e:
        print("OpenAI API Error:", e)
        return jsonify(success=False, message="OpenAI API Error"), 500

    cover_letter_content = completion.choices[0].message.content
    base_filename = f"{company_name}_cv.docx"
    filename = base_filename
    count = 1

    # Check if "Generated_CVs" folder exists, if not, create it
    folder_name = "Generated_CVs"
    if not os.path.exists(folder_name):
        os.mkdir(folder_name)

    # Ensure unique filename within the "Generated_CVs" folder
    while os.path.isfile(os.path.join(folder_name, filename)):
        filename = f"{company_name}_cv({count}).docx"
        count += 1

    # Create the DOCX in memory
    doc = Document()
    doc.add_paragraph(cover_letter_content)
    mem_stream = io.BytesIO()
    doc.save(mem_stream)

    # Reset stream position
    mem_stream.seek(0)

    # Increment user's counter
    user_id_decode = str(user_id_decode)  # Ensure the key is a string
    user_counters[user_id_decode] = user_counters.get(user_id_decode, 0) + 1
    print(f"User {user_id_decode} has generated {user_counters[user_id_decode]} cover letters.")

    # Send the in-memory DOCX as a file
    return send_file(mem_stream, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')



# resume generation
@app.route("/generate-resume", methods=["POST"])
def generate_resume():
    token = request.headers.get("Authorization")
    try:
        data = serializer.loads(token, salt="password-reset", max_age=9600)
    except Exception as e:
        print(f"Token deserialization error: {e}")
        return jsonify(success=False, message="Invalid or expired token"), 401


    data = request.get_json()
    api_key = data.get("apiKey")
    openai.api_key = api_key

    job_description = data.get("Job-Description", "")
    user_id = data.get("user_id")

    print(user_id)
    print(job_description)

    decoded_data = serializer.loads(user_id, salt="password-reset", max_age=9600)

    user_id_decode = decoded_data["user"]
    print("User ID:", user_id_decode)

    # Construct the path to the user's folder and the resume inside
    user_folder_path = os.path.join(UPLOAD_FOLDER, str(user_id_decode))
    resume_path = os.path.join(user_folder_path, "current_resume.txt")

    try:
        with open(resume_path, "r") as file:
            resume = file.read()
    except FileNotFoundError:
        return jsonify(success=False, message="Resume file not found."), 404
    
    print(resume)  

    try:
        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system",
                    "content": "You are a helpful assistant that will reword a resume to better align it with a job description. Your goal is to highlight the candidate's skills and experiences, making them match the job's requirements as closely as possible.",
                },
                {
                    "role": "user",
                    "content": f"Given this original resume: {resume}, and this job description: {job_description}, please reword the resume to better fit the job requirements.",
                },
            ],
            temperature=1.3,
            top_p=0.9,
            max_tokens=1000,
            frequency_penalty=0.5,
            presence_penalty=0.5,
        )
    except openai.error.OpenAIError as e:
        print("OpenAI API Error:", e)
        return jsonify(success=False, message="OpenAI API Error"), 500

    reworded_resume_content = completion.choices[0].message.content
    base_filename = "reworded_resume.docx"
    filename = base_filename
    count = 1

    # Check if "Generated_Resumes" folder exists, if not, create it
    folder_name = "Generated_Resumes"
    if not os.path.exists(folder_name):
        os.mkdir(folder_name)

    # Ensure unique filename within the "Generated_Resumes" folder
    while os.path.isfile(os.path.join(folder_name, filename)):
        filename = f"reworded_resume({count}).docx"
        count += 1

    # Create the DOCX in memory
    doc = Document()
    doc.add_paragraph(reworded_resume_content)
    mem_stream = io.BytesIO()
    doc.save(mem_stream)

    # Reset stream position
    mem_stream.seek(0)

    # Increment user's counter
    user_id_decode = str(user_id_decode)  # Ensure the key is a string
    user_counters[user_id_decode] = user_counters.get(user_id_decode, 0) + 1
    print(f"User {user_id_decode} has generated {user_counters[user_id_decode]} cover letters.")

    # Send the in-memory DOCX as a file
    return send_file(mem_stream, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')



# user registration
@app.route("/request-reset-password", methods=["POST"])
def request_reset_password():
    data = request.get_json()
    email = data.get("email")
    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify(success=False, message="User not found"), 404

    token = serializer.dumps({"user": user.user_id}, salt="password-reset")
    reset_url = url_for("reset_password_with_token", token=token, _external=True)
    msg = Message(
        "Password Reset Request", sender="lynktools@gmail.com", recipients=[email]
    )
    msg.body = f"To reset your password, click on the following link: {reset_url}"
    print(msg.body)
    try:
        mail.send(msg)
    except Exception as e:
        print("Error sending email:", e)
        return jsonify(success=False, message="Error sending email."), 500


    return jsonify(success=True, message="Password reset email has been sent."), 200


@app.route("/reset-password/<token>", methods=["GET", "POST"])
def reset_password_with_token(token):
    try:
        # This will raise an exception if the token is invalid or has expired
        data = serializer.loads(token, salt="password-reset", max_age=9600)
    except:
        return jsonify(success=False, message="Invalid or expired token"), 401

    user_id = data["user"]
    user = User.query.get(user_id)

    if not user:
        return jsonify(success=False, message="User not found"), 404

    # Handle the GET request by rendering the password reset form
    if request.method == "GET":
        frontend_url = "http://127.0.0.1:3000/reset-password/{token}".format(token=token)
        return redirect(frontend_url)


    # Handle the POST request
    elif request.method == "POST":
        # Fetch the new password from the form data
        new_password_data = request.get_json()
        new_password = new_password_data.get("newPassword")
        hashed_password = generate_password_hash(new_password, method="sha256")
        user.password = hashed_password
        db.session.commit()

        return jsonify(success=True, message="Password reset successful"), 200
    
@app.route("/verify-reset-token/<token>", methods=["GET"])
def verify_reset_token(token):
    try:
        # This will raise an exception if the token is invalid or has expired
        data = serializer.loads(token, salt="password-reset", max_age=9600)
    except:
        return jsonify(success=False, message="Invalid or expired token"), 401

    user_id = data["user"]
    user = User.query.get(user_id)

    if not user:
        return jsonify(success=False, message="User not found"), 404

    return jsonify(success=True, message="Valid token"), 200



@app.route("/upload-resume", methods=["POST"])
def upload_resume():
    # Retrieve the token from formData
    token = request.form.get("user_id")
    print(token)

    # Check if the token is provided
    if not token:
        return jsonify(success=False, message="Token not provided"), 421

    # Deserialize the token to extract user_id
    try:
        data = serializer.loads(token, salt="password-reset", max_age=9600)
        user_id = data["user"]
        print(user_id)
    except Exception as e:
        print(f"Deserialization error: {e}")
        return jsonify(success=False, message="Invalid or expired token"), 420


    # Check if file is present in the request
    if "resume" not in request.files:
        return jsonify(success=False, message="No file part"), 400

    file = request.files["resume"]

    # Check if a filename is provided
    if file.filename == "":
        return jsonify(success=False, message="No selected file"), 400

    # Check for allowed file extensions
    ALLOWED_EXTENSIONS = ["pdf", "docx"]
    if file.filename.rsplit(".", 1)[1].lower() not in ALLOWED_EXTENSIONS:
        return jsonify(success=False, message="File type not allowed"), 400

    # Ensure file is not too large
    if len(file.read()) > 5 * 1024 * 1024:
        return jsonify(success=False, message="File size exceeds the limit"), 400

    # Reset file pointer after reading
    file.seek(0)
    file_type = os.path.splitext(file.filename)[1][1:]

    # Convert file to txt
    try:
        txt_content = convert_to_txt(file, file_type)
    except ValueError:
        return jsonify(success=False, message="Unsupported file type"), 400
    except Exception as e:
        return jsonify(success=False, message=str(e)), 500

    # Store in database
    resume = Resume(user_id=user_id, content=txt_content)
    db.session.add(resume)
    db.session.commit()

    # Store the text content in the user's folder
    user_folder = os.path.join(UPLOAD_FOLDER, str(user_id))
    if not os.path.exists(user_folder):
        os.mkdir(user_folder)

    filename = os.path.join(user_folder, "current_resume.txt")
    with open(filename, "w", encoding="utf-8") as txt_file:
        txt_file.write(txt_content)

    return (
        jsonify(success=True, message="File uploaded and converted successfully"),
        200,
    )


# user authentication
@app.route("/authenticate", methods=["POST"])
def authenticate():
    data = request.get_json()
    print(data)
    action = data.get("action")
    email = data.get("email")
    password = data.get("password")

    if action == "register":
        name = data.get("name")
        hashed_password = encrypt_password(password)
        print(hashed_password)

        user = User.query.filter_by(email=email).first()
        if user:
            return jsonify(success=False, message="User already exists"), 400

        user_role = Role.query.filter_by(name="user").first()
        if not user_role:
            user_role = Role(name="user", description="Regular user")
            db.session.add(user_role)
            db.session.commit()

        new_user = User(email=email, password=hashed_password)
        new_user.roles.append(user_role)
        db.session.add(new_user)
        db.session.commit()

        return jsonify(success=True, message="User registered successfully"), 200

    elif action == "login":
        user = User.query.filter_by(email=email).first()
    if not user or not verify_password(user.password, password):
        return jsonify(success=False, message="Invalid email or password"), 401

    # session["user_id"] = user.user_id

    # Generate token or handle login as needed
    token = serializer.dumps({"user": user.user_id}, salt="password-reset")
    print(token)
    # print(user)
    return jsonify(success=True, token=token), 200


# admin page
@app.route("/admin")
@roles_required("admin")
@login_required
def admin():
    return "Admin Page"
