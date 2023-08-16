from werkzeug.security import generate_password_hash as encrypt_password
from werkzeug.security import check_password_hash as verify_password
from flask_security.forms import RegisterForm, LoginForm
from wtforms import StringField
from wtforms.validators import DataRequired
from flask import Flask, request, send_file, jsonify, url_for
from dotenv import load_dotenv
import openai
from docx import Document
import os
from flask_cors import CORS
import PyPDF2
from itsdangerous import URLSafeTimedSerializer
from flask_security import (
    Security,
    SQLAlchemyUserDatastore,
    UserMixin,
    RoleMixin,
    roles_required,
    login_required,
)
from flask_sqlalchemy import SQLAlchemy
import secrets
import uuid
from flask_session import Session
import redis
import logging
from flask_mail import Mail, Message
from flask_mail import Mail as FlaskMail, Message
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail


app = Flask(__name__)
load_dotenv()
# Session configurations
DB_USERNAME = os.getenv("DB_USERNAME")
DB_PASSWORD = os.getenv("DB_PASSWORD")
DB_HOST = os.getenv("DB_HOST")
DB_PORT = os.getenv("DB_PORT")
DB_NAME = os.getenv("DB_NAME")

app.config[
    "SQLALCHEMY_DATABASE_URI"
] = f"postgresql://{DB_USERNAME}:{DB_PASSWORD}@{DB_HOST}:{DB_PORT}/{DB_NAME}"


app.config["SESSION_TYPE"] = "redis"
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_USE_SIGNER"] = True
app.config["SESSION_KEY_PREFIX"] = "your_app:"
app.config["SESSION_REDIS"] = redis.StrictRedis(
    host=os.getenv("REDIS_HOST", "localhost"), port=os.getenv("REDIS_PORT", 6379), db=0
)
sess = Session()
sess.init_app(app)

app.config["YOUR_SENDGRID_API_KEY"] = os.getenv("SENDGRID_API_KEY")

CORS(
    app,
    resources={r"/*": {"origins": ["https://www.linkedin.com"]}},
    supports_credentials=True,
)

# Database configuration
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)

SECRET_KEY = os.getenv("SECRET_KEY", default=secrets.token_urlsafe(16))
app.config["SECRET_KEY"] = SECRET_KEY
serializer = URLSafeTimedSerializer(app.config["SECRET_KEY"])

app.config["SECURITY_PASSWORD_SALT"] = os.getenv(
    "SECURITY_PASSWORD_SALT", default="your_random_salt"
)
app.config["SECURITY_REGISTERABLE"] = True
app.config["SECURITY_RECOVERABLE"] = True
app.config["SECURITY_PASSWORD_SALT"] = os.getenv(
    "SECURITY_PASSWORD_SALT", default="your_random_salt"
)

app.config["MAIL_SERVER"] = "smtp.sendgrid.net"
app.config["MAIL_PORT"] = 587  # 465 for TLS
app.config["MAIL_USERNAME"] = "apikey"
app.config["MAIL_PASSWORD"] = os.environ.get("MAIL_PASSWORD")
app.config["MAIL_USE_TLS"] = True
app.config["MAIL_USE_SSL"] = False


mail = FlaskMail(app)

UPLOAD_FOLDER = os.path.dirname(os.path.abspath(__file__))

# Configuration options
app.config["SECURITY_REGISTERABLE"] = True
app.config["SECURITY_RECOVERABLE"] = True
app.config["SECURITY_TRACKABLE"] = True

logging.basicConfig(level=logging.INFO)


class CoverLetter(db.Model):
    cover_letter_id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.user_id"), nullable=False)
    company_name = db.Column(db.String(100))
    job_listing = db.Column(db.Text)
    recruiter = db.Column(db.String(100))
    date = db.Column(db.DateTime, default=db.func.current())
    file_path = db.Column(db.String(255))
    user = db.relationship("User", backref="cover_letters")


class Resume(db.Model):
    resume_id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.user_id"), nullable=False)
    content = db.Column(db.Text)
    upload_date = db.Column(db.DateTime, default=db.func.current())
    user = db.relationship("User", backref="resumes")


class UsageStatistic(db.Model):
    stat_id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.user_id"))
    action = db.Column(db.String(50))
    date = db.Column(db.DateTime, default=db.func.current())
    user = db.relationship("User", backref="usage_statistics")


# Define roles
class Role(db.Model, RoleMixin):
    id = db.Column(db.Integer(), primary_key=True)
    name = db.Column(db.String(80), unique=True)
    description = db.Column(db.String(255))


roles_users = db.Table(
    "roles_users",
    db.Column("user_id", db.Integer(), db.ForeignKey("user.user_id")),
    db.Column("role_id", db.Integer(), db.ForeignKey("role.id")),
)


class User(db.Model, UserMixin):
    user_id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(255))
    password = db.Column(db.String(255), nullable=False)
    email = db.Column(db.String(150), unique=True, nullable=False)
    fs_uniquifier = db.Column(db.String(255), unique=True)  # Add this line
    roles = db.relationship(
        "Role", secondary=roles_users, backref=db.backref("users", lazy="dynamic")
    )


with app.app_context():
    db.create_all()

# Define hashed_password
password = "supersecretpassword"
hashed_password = encrypt_password(password)

# Setup Flask-Security
user_datastore = SQLAlchemyUserDatastore(db, User, Role)


# Custom registration form
class ExtendedRegisterForm(RegisterForm):
    username = StringField("Username", [DataRequired()])


# Configuring Flask-Security with the custom registration form
app.config["SECURITY_REGISTER_USER_TEMPLATE"] = "security/register_user.html"
app.config["SECURITY_REGISTERABLE"] = True
app.config["SECURITY_CONFIRMABLE"] = True
app.config["SECURITY_RECOVERABLE"] = True
app.config["SECURITY_REGISTER_USER_TEMPLATE"] = "security/register_user.html"
app.config["SECURITY_LOGIN_USER_TEMPLATE"] = "security/login_user.html"
security = Security(app, user_datastore, register_form=ExtendedRegisterForm)


def init_security(app, user_datastore):
    return Security(app, user_datastore)


def is_api_key_valid(api_key):
    openai.api_key = api_key
    try:
        response = openai.Completion.create(
            engine="davinci", prompt="This is a test.", max_tokens=5
        )
    except:
        return False
    else:
        return True


def convert_to_txt(file, file_type):
    if file_type == "docx":
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif file_type == "pdf":
        reader = PyPDF2.PdfReader(file)
        return "\n".join(
            [reader.pages[i].extract_text() for i in range(len(reader.pages))]
        )
    else:
        raise ValueError("Unsupported file type")


@app.route("/apiverify", methods=["POST"])
def apiverify():
    data = request.get_json()
    api_key = data.get("apiKey")

    try:
        if is_api_key_valid(api_key):
            token = serializer.dumps({"user": "YOUR_USER_IDENTIFIER"})
            return jsonify(success=True, token=token)
        else:
            raise Exception("Invalid API key")
    except Exception as e:
        logging.error(str(e))
        return jsonify(success=False, message="Invalid API key"), 401


@app.route("/cover-letter", methods=["POST", "GET", "PUT", "DELETE"])
def listen():
    token = request.headers.get("Authorization")
    try:
        data = serializer.loads(token, max_age=3600)
    except:
        return jsonify(success=False, message="Invalid or expired token"), 401

    data = request.get_json()
    api_key = data.get("apiKey")
    openai.api_key = api_key

    company_name = data.get("Company-name", "")
    job_listing = data.get("Job-Listing", "")
    recruiter = data.get("Recruiter", "")
    date = data.get("Date", "")

    try:
        resume = open("current_resume.txt", "r").read()
    except FileNotFoundError:
        return "Resume file not found.", 404

    try:
        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system",
                    "content": "You are a helpful assistant that will craft a tailored cover letter using a resume and a job listing. Your goal is to create a compelling cover letter that showcases the candidate's skills and experiences, aligning them with the job's requirements.",
                },
                {
                    "role": "user",
                    "content": f"Using this resume, {resume}, and this job listing, {job_listing}, craft a cover letter that doesn't include addresses but highlights the candidate's fit for the role. Ensure it includes the candidate's name, email, phone number, and LinkedIn profile. Also, only include the company name {company_name}, followed by recruiter's name {recruiter} and today's date {date}. No place holder text is allowed, if recruiters name is not found use 'Dear Hiring Manager.' ",
                },
            ],
            temperature=1.3,
            top_p=0.9,
            max_tokens=700,
            frequency_penalty=0.5,
            presence_penalty=0.5,
        )
    except openai.error.OpenAIError as e:
        print("OpenAI API Error:", e)
        return jsonify(success=False, message="OpenAI API Error"), 500

    cover_letter_content = completion.choices[0].message.content
    base_filename = f"{company_name}_cv.docx"
    filename = base_filename
    count = 1

    # Check if "Generated_CVs" folder exists, if not, create it
    folder_name = "Generated_CVs"
    if not os.path.exists(folder_name):
        os.mkdir(folder_name)

    # Ensure unique filename within the "Generated_CVs" folder
    while os.path.isfile(os.path.join(folder_name, filename)):
        filename = f"{company_name}_cv({count}).docx"
        count += 1

    # Create the document and save it inside "Generated_CVs" folder
    doc = Document()
    doc.add_paragraph(cover_letter_content)
    full_path = os.path.join(folder_name, filename)
    doc.save(full_path)

    return send_file(full_path, as_attachment=True, download_name=filename)


@app.route("/request-reset-password", methods=["POST"])
def request_reset_password():
    data = request.get_json()
    email = data.get("email")
    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify(success=False, message="User not found"), 404

    token = serializer.dumps({"user": user.user_id}, salt="password-reset")
    reset_url = url_for("request_reset_password", token=token, _external=True)
    msg = Message(
        "Password Reset Request", sender="lynktools@gmail.com", recipients=[email]
    )
    msg.body = f"To reset your password, click on the following link: {reset_url}"
    mail.send(msg)

    return jsonify(success=True, message="Password reset email has been sent."), 200


from werkzeug.security import generate_password_hash


@app.route("/reset-password/<token>", methods=["POST"])
def reset_password_with_token(token):
    try:
        # This will raise an exception if the token is invalid or has expired
        data = serializer.loads(token, salt="password-reset", max_age=3600)
    except:
        return jsonify(success=False, message="Invalid or expired token"), 401

    user_id = data["user"]
    user = User.query.get(user_id)

    if not user:
        return jsonify(success=False, message="User not found"), 404

    # Here, you can fetch the user using the user_id and update their password
    new_password = request.json.get("newPassword")
    hashed_password = generate_password_hash(new_password, method="sha256")
    user.password = hashed_password
    db.session.commit()

    return jsonify(success=True, message="Password reset successful"), 200


@app.route("/upload-resume", methods=["POST"])
def upload_resume():
    if "resume" not in request.files:
        return jsonify(success=False, message="No file part"), 400

    # need to add user_id to the request
    user_id = request.form.get("user_id")
    file = request.files["resume"]

    if file.filename == "":
        return jsonify(success=False, message="No selected file"), 400

    # Allowed file extensions
    ALLOWED_EXTENSIONS = ["pdf", "docx"]
    if file and file.filename.rsplit(".", 1)[1].lower() not in ALLOWED_EXTENSIONS:
        return jsonify(success=False, message="File type not allowed"), 400

    if file and (len(file.read()) <= 5 * 1024 * 1024):
        file.seek(0)
        file_type = os.path.splitext(file.filename)[1][1:]

    try:
        txt_content = convert_to_txt(file, file_type)
    except Exception as e:
        return jsonify(success=False, message=str(e)), 500
    except ValueError:
        return jsonify(success=False, message="Unsupported file type"), 400

    resume = Resume(user_id=user_id, content=txt_content)

    # db stuff
    db.session.add(resume)
    db.session.commit()

    user_folder = os.path.join(UPLOAD_FOLDER, str(user_id))

    if not os.path.exists(user_folder):
        os.mkdir(user_folder)

    filename = os.path.join(user_folder, "current_resume.txt")
    with open(filename, "w", encoding="utf-8") as txt_file:
        txt_file.write(txt_content)

    return (
        jsonify(success=True, message="File uploaded and converted successfully"),
        200,
    )


@app.route("/authenticate", methods=["POST"])
def authenticate():
    data = request.get_json()
    print(data)
    action = data.get("action")
    email = data.get("email")
    password = data.get("password")

    if action == "register":
        name = data.get("name")
        hashed_password = encrypt_password(password)
        print(hashed_password)

        user = User.query.filter_by(email=email).first()
        if user:
            return jsonify(success=False, message="User already exists"), 400

        user_role = Role.query.filter_by(name="user").first()
        if not user_role:
            user_role = Role(name="user", description="Regular user")
            db.session.add(user_role)
            db.session.commit()

        new_user = User(email=email, password=hashed_password)
        new_user.roles.append(user_role)
        db.session.add(new_user)
        db.session.commit()

        return jsonify(success=True, message="User registered successfully"), 200

    elif action == "login":
        user = User.query.filter_by(email=email).first()
    if not user or not verify_password(user.password, password):
        return jsonify(success=False, message="Invalid email or password"), 401

    # Generate token or handle login as needed
    token = serializer.dumps({"user": user.user_id}, salt="password-reset")
    return jsonify(success=True, token=token), 200



@app.route("/admin")
@roles_required("admin")
@login_required
def admin():
    return "Admin Page"


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=3000, debug=True)
